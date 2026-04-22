import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_readme_docx():
    doc = Document()

    # Заголовок проекта
    title = doc.add_heading('Team Finder', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Краткое описание
    p = doc.add_paragraph()
    p.add_run('Веб-приложение для эффективного поиска участников и формирования команд для различных проектов.').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Секция: О проекте
    doc.add_heading('Описание проекта', level=1)
    doc.add_paragraph(
        "Team Finder — это платформа, объединяющая специалистов из разных областей "
        "(разработчиков, дизайнеров, менеджеров) для совместной работы. Проект решает "
        "проблему поиска единомышленников для реализации стартапов, учебных или open-source инициатив."
    )

    # Секция: Стек технологий
    doc.add_heading('Стек технологий', level=1)
    tech_stack = [
        ('Язык программирования', 'Python 3.10+'),
        ('Фреймворк', 'Django'),
        ('База данных', 'PostgreSQL 16'),
        ('Инфраструктура', 'Docker / Docker Compose'),
        ('Веб-сервер', 'Nginx, Gunicorn')
    ]
    for key, value in tech_stack:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{key}: ").bold = True
        p.add_run(value)

    # Секция: Установка и запуск
    doc.add_heading('Установка и запуск', level=1)

    doc.add_heading('1. Клонирование репозитория', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'git clone git@github.com:YarkosEnergos/team-finder-ad.git\ncd team-finder-ad')
    run.font.name = 'Courier New'

    doc.add_heading('2. Настройка окружения', level=2)
    doc.add_paragraph(
        'Создайте файл .env в корне проекта и добавьте следующие настройки:')
    env_code = (
        "POSTGRES_DB=teamfinder\n"
        "POSTGRES_USER=postgres\n"
        "POSTGRES_PASSWORD=postgres\n"
        "POSTGRES_HOST=db\n"
        "POSTGRES_PORT=5432\n"
        "SECRET_KEY=your_secret_key\n"
        "DEBUG=1\n"
        "ALLOWED_HOSTS=localhost,127.0.0.1"
    )
    p = doc.add_paragraph()
    run = p.add_run(env_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('3. Развертывание через Docker', level=2)
    docker_cmds = [
        'docker compose up -d --build',
        'docker compose exec backend python manage.py migrate',
        'docker compose exec backend python manage.py createsuperuser',
        'docker compose exec backend python manage.py collectstatic --noinput'
    ]
    for cmd in docker_cmds:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(cmd)
        run.font.name = 'Courier New'

    # Секция: Доступ
    doc.add_heading('Доступ к приложению', level=1)
    doc.add_paragraph('После запуска сервисы будут доступны по адресам:')
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Сайт: http://localhost:8000')
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Админка: http://localhost:8000/admin')

    # Секция: Полезные команды
    doc.add_heading('Команды для управления', level=1)
    mgmt_cmds = [
        'docker compose down (остановить проект)',
        'docker compose down -v (остановить и удалить данные БД)',
        'docker compose logs -f (просмотр логов)',
        'docker compose exec backend bash (доступ к терминалу контейнера)'
    ]
    for cmd in mgmt_cmds:
        doc.add_paragraph(cmd, style='List Bullet')

    # Секция: Автор
    doc.add_heading('Об авторе', level=1)
    doc.add_paragraph('Автор проекта: YarkosEnergos')
    p = doc.add_paragraph()
    p.add_run('GitHub: https://github.com/YarkosEnergos\nEmail: contact@example.com')

    # Сохранение
    file_path = 'README-Team-Finder.docx'
    doc.save(file_path)
    return file_path


create_readme_docx()
